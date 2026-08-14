"""
Tests for .docx content that lives outside the ordinary paragraph model,
and for the residual scan that re-checks the finished file.

Every test here corresponds to a leak that was verified to exist before the
feature was written -- e.g. a tracked-change insertion really does report
zero runs to `python-docx`, so its text really was invisible to every
detector.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest
from docx import Document

from app.core.detectors.base import DetectionConfig
from app.core.document_io import redact_file
from app.core.redactor import Redactor
from conftest import (all_xml_text, build_docx_with_hidden_content,
                      build_docx_with_revisions)


def _redactor():
    return Redactor(config=DetectionConfig())


class TestTrackedChangesAndHyperlinks:
    """
    `paragraph.runs` returns only `<w:r>` elements that are DIRECT children
    of `<w:p>`. Word nests runs inside `<w:ins>`, `<w:del>` and
    `<w:hyperlink>`, so all three were invisible to the detectors.
    """

    def test_insertion_is_invisible_before_flattening(self, tmp_path):
        # Guards the premise of the whole feature: if python-docx ever
        # starts returning these runs, `flatten_revisions` can be revisited.
        path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        doc = Document(path)
        insertion = [p for p in doc.paragraphs if "Inserted" in p.text or len(p.runs) == 0]
        assert any(len(p.runs) == 0 for p in insertion), \
            "expected a tracked-change paragraph to expose no runs"

    def test_tracked_insertion_is_redacted(self, tmp_path):
        in_path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        out_path = str(tmp_path / "out.docx")
        redact_file(in_path, out_path, _redactor())
        xml = all_xml_text(out_path)
        assert "Priya Sharma" not in xml
        assert "priya.sharma@example.org" not in xml

    def test_tracked_deletion_is_removed_entirely(self, tmp_path):
        # A deletion is still physically in the file and visible the moment
        # someone turns on "Show Markup", so it must not survive at all.
        in_path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        out_path = str(tmp_path / "out.docx")
        redact_file(in_path, out_path, _redactor())
        xml = all_xml_text(out_path)
        assert "Angela Brooks" not in xml
        assert "angela.brooks@example.org" not in xml

    def test_hyperlink_display_text_is_redacted(self, tmp_path):
        in_path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        out_path = str(tmp_path / "out.docx")
        redact_file(in_path, out_path, _redactor())
        assert "chitra.raste@example.net" not in all_xml_text(out_path)


class TestFootnotesEndnotesCommentsProperties:
    """Surfaces with either no python-docx API at all (footnotes, endnotes,
    custom properties) or an API that was never wired in (comments)."""

    @staticmethod
    def _build(tmp_path, **overrides):
        params = dict(
            footnote_text="See Manisha Shukla, manisha.shukla@example.io",
            comment_text="Verify with Sarthak Malvadkar at sarthak.m@example.com",
            comment_author="Jane Reviewer",
            custom_property_text="Tushar Wakhele",
        )
        params.update(overrides)
        return build_docx_with_hidden_content(str(tmp_path / "hidden.docx"), **params)

    def test_footnote_and_endnote_text_is_redacted(self, tmp_path):
        out_path = str(tmp_path / "out.docx")
        redact_file(self._build(tmp_path), out_path, _redactor())
        xml = all_xml_text(out_path)
        assert "Manisha Shukla" not in xml
        assert "manisha.shukla@example.io" not in xml

    def test_comment_text_is_redacted(self, tmp_path):
        out_path = str(tmp_path / "out.docx")
        redact_file(self._build(tmp_path), out_path, _redactor())
        xml = all_xml_text(out_path)
        assert "Sarthak Malvadkar" not in xml
        assert "sarthak.m@example.com" not in xml

    def test_comment_author_metadata_is_scrubbed(self, tmp_path):
        # The author is a real name stored in the comment's metadata,
        # independent of anything the comment text says.
        out_path = str(tmp_path / "out.docx")
        result = redact_file(self._build(tmp_path), out_path, _redactor())
        assert "Jane Reviewer" not in all_xml_text(out_path)
        # ...and it is recorded in the audit trail like any other redaction.
        assert any(r["source"] == "metadata" for r in result.replacements)

    def test_custom_document_property_is_redacted(self, tmp_path):
        out_path = str(tmp_path / "out.docx")
        result = redact_file(self._build(tmp_path), out_path, _redactor())
        xml = all_xml_text(out_path)
        assert "Tushar Wakhele" not in xml
        # A property with no PII in it is left exactly as it was.
        assert "Legal" in xml
        assert any("custom document propert" in w for w in result.warnings)

    def test_core_properties_author_is_scrubbed(self, tmp_path):
        out_path = str(tmp_path / "out.docx")
        redact_file(self._build(tmp_path), out_path, _redactor())
        assert "Real Author Name" not in all_xml_text(out_path)


class TestUnredactableContentWarnings:
    """The tool cannot read text inside an image, so it says so rather than
    reporting a clean result that quietly isn't."""

    def test_embedded_image_produces_a_warning(self, tmp_path):
        path = build_docx_with_hidden_content(
            str(tmp_path / "img.docx"),
            footnote_text="nothing sensitive",
            comment_text="nothing sensitive",
            comment_author="A B",
            custom_property_text="nothing sensitive",
            include_picture=True,
        )
        result = redact_file(path, str(tmp_path / "out.docx"), _redactor())
        assert any("image" in w.lower() and "OCR" in w for w in result.warnings)

    def test_no_image_means_no_image_warning(self, tmp_path):
        path = build_docx_with_hidden_content(
            str(tmp_path / "noimg.docx"),
            footnote_text="nothing sensitive",
            comment_text="nothing sensitive",
            comment_author="A B",
            custom_property_text="nothing sensitive",
            include_picture=False,
        )
        result = redact_file(path, str(tmp_path / "out.docx"), _redactor())
        assert not any("image" in w.lower() for w in result.warnings)


class TestResidualScan:
    """The tool re-opens the finished file and checks its own work."""

    def test_clean_document_reports_clean(self, tmp_path):
        path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        result = redact_file(path, str(tmp_path / "out.docx"), _redactor())
        assert result.residual is not None
        assert result.residual.clean
        assert result.residual.leaked_originals == []

    def test_injected_leak_is_caught(self, tmp_path):
        from app.core.verification import verify_docx

        in_path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        out_path = str(tmp_path / "out.docx")
        result = redact_file(in_path, out_path, _redactor())

        # Simulate the pipeline believing it redacted a name that is, in
        # fact, still sitting in the output.
        doc = Document(out_path)
        doc.add_paragraph("Leftover: Kushal Subbayya Hegde, kushal@realbank.com")
        leaked_path = str(tmp_path / "leaked.docx")
        doc.save(leaked_path)

        claimed = result.replacements + [{
            "type": "PERSON", "original": "Kushal Subbayya Hegde",
            "fake": "Someone Else", "confidence": 1.0, "source": "ner",
        }]
        report = verify_docx(leaked_path, claimed, _redactor())

        assert not report.clean
        assert any(f.text == "Kushal Subbayya Hegde" for f in report.leaked_originals)

    def test_summary_never_exposes_the_matched_text(self, tmp_path):
        # The summary is what goes over the API. A redaction service must
        # not echo potentially-sensitive strings back in its own response.
        from app.core.verification import verify_docx

        in_path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        out_path = str(tmp_path / "out.docx")
        result = redact_file(in_path, out_path, _redactor())
        summary = verify_docx(out_path, result.replacements, _redactor()).summary()
        assert set(summary) == {
            "clean", "leaked_original_count", "leaked_original_types",
            "unexpected_match_count", "unexpected_match_types",
        }
        assert all(not isinstance(v, str) or v in ("",) for k, v in summary.items()
                   if k.endswith("count"))

    def test_verification_can_be_skipped(self, tmp_path):
        path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        result = redact_file(path, str(tmp_path / "out.docx"), _redactor(), verify=False)
        assert result.residual is None


class TestConfidenceTiers:
    def test_every_replacement_carries_a_tier(self, tmp_path):
        path = build_docx_with_revisions(str(tmp_path / "rev.docx"))
        result = redact_file(path, str(tmp_path / "out.docx"), _redactor())
        assert result.replacements
        assert all(r["confidence_tier"] in ("high", "medium", "needs_review")
                   for r in result.replacements)

    def test_tier_thresholds(self):
        from app.core.detectors.base import confidence_tier
        assert confidence_tier(0.99) == "high"      # validated email
        assert confidence_tier(0.90) == "high"      # boundary
        assert confidence_tier(0.85) == "medium"    # ordinary NER
        assert confidence_tier(0.80) == "medium"    # boundary
        assert confidence_tier(0.75) == "needs_review"  # consistency sweep
        assert confidence_tier(0.60) == "needs_review"  # Aadhaar, no context


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
