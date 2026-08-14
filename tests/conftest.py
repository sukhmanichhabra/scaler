"""
Shared fixture builders for .docx surfaces that `python-docx` cannot create
through its normal API.

Footnotes, endnotes and custom document properties have no `python-docx`
object model at all, so the only way to get a test document containing them
is to write the parts into the .docx zip directly -- which is exactly what
Word itself does. The XML here is minimal but structurally valid: each part
is declared in `[Content_Types].xml` and related from the correct `.rels`
file, or `python-docx` won't see it when the document is reopened.
"""
from __future__ import annotations

import zipfile

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

_FOOTNOTES_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:footnote w:type="separator" w:id="-1"><w:p/></w:footnote>
<w:footnote w:type="continuationSeparator" w:id="0"><w:p/></w:footnote>
<w:footnote w:id="1"><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:footnote>
</w:footnotes>"""

_CUSTOM_PROPS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" \
xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="2" name="Prepared For">\
<vt:lpwstr>{text}</vt:lpwstr></property>
<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="3" name="Department">\
<vt:lpwstr>Legal</vt:lpwstr></property>
</Properties>"""

_PICTURE_RUN = (
    '<w:r><w:drawing>'
    '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
    '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
    '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
    '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
    '<pic:nvPicPr><pic:cNvPr id="1" name="img"/></pic:nvPicPr></pic:pic>'
    '</a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
)


def build_docx_with_revisions(path: str) -> str:
    """
    A document whose PII sits in places `paragraph.runs` does not reach:
    a tracked-change insertion, a tracked-change deletion, and a hyperlink.

    Verified behaviour without `flatten_revisions`: the insertion and
    deletion paragraphs both report `len(paragraph.runs) == 0`.
    """
    doc = Document()
    doc.add_paragraph("Body mentions Rohan Dey at rohan.dey@gmail.com.")

    doc.element.body.append(parse_xml(
        '<w:p %s><w:ins w:id="1" w:author="Reviewer" w:date="2024-01-01T00:00:00Z">'
        "<w:r><w:t>Inserted: Priya Sharma priya.sharma@example.org</w:t></w:r></w:ins></w:p>"
        % nsdecls("w")))

    doc.element.body.append(parse_xml(
        '<w:p %s><w:del w:id="2" w:author="Reviewer" w:date="2024-01-01T00:00:00Z">'
        "<w:r><w:delText>Deleted: Angela Brooks angela.brooks@example.org</w:delText>"
        "</w:r></w:del></w:p>" % nsdecls("w")))

    doc.element.body.append(parse_xml(
        '<w:p %s><w:r><w:t>Link: </w:t></w:r>'
        '<w:hyperlink r:id="rId9"><w:r><w:t>chitra.raste@example.net</w:t></w:r></w:hyperlink></w:p>'
        % nsdecls("w", "r")))

    doc.save(path)
    return path


def build_docx_with_hidden_content(path: str, *, footnote_text: str,
                                    comment_text: str, comment_author: str,
                                    custom_property_text: str,
                                    include_picture: bool = True) -> str:
    """
    A document exercising every hidden-content surface at once: a footnote,
    an endnote, a comment (with a real author name in its metadata), a
    custom document property, and optionally an embedded picture.
    """
    doc = Document()
    doc.core_properties.author = "Real Author Name"
    doc.add_paragraph("Ordinary body text.")
    doc.part.comments.add_comment(text=comment_text, author=comment_author, initials="JR")

    base_path = path + ".base"
    doc.save(base_path)

    footnotes = _FOOTNOTES_XML.format(text=footnote_text)
    endnotes = footnotes.replace("footnote", "endnote")
    custom_props = _CUSTOM_PROPS_XML.format(text=custom_property_text)

    body_suffix = b'<w:p><w:r><w:t>Ref</w:t><w:footnoteReference w:id="1"/>'
    body_suffix += b'<w:endnoteReference w:id="1"/></w:r>'
    if include_picture:
        body_suffix += _PICTURE_RUN.encode()
    body_suffix += b"</w:p></w:body>"

    with zipfile.ZipFile(base_path) as zin, \
            zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = data.replace(b"</w:body>", body_suffix)
            elif item.filename == "[Content_Types].xml":
                data = data.replace(b"</Types>", (
                    b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.'
                    b'openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                    b'<Override PartName="/word/endnotes.xml" ContentType="application/vnd.'
                    b'openxmlformats-officedocument.wordprocessingml.endnotes+xml"/>'
                    b'<Override PartName="/docProps/custom.xml" ContentType="application/vnd.'
                    b'openxmlformats-officedocument.custom-properties+xml"/></Types>'))
            elif item.filename == "word/_rels/document.xml.rels":
                data = data.replace(b"</Relationships>", (
                    b'<Relationship Id="rIdFn1" Type="http://schemas.openxmlformats.org/'
                    b'officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                    b'<Relationship Id="rIdEn1" Type="http://schemas.openxmlformats.org/'
                    b'officeDocument/2006/relationships/endnotes" Target="endnotes.xml"/>'
                    b"</Relationships>"))
            elif item.filename == "_rels/.rels":
                data = data.replace(b"</Relationships>", (
                    b'<Relationship Id="rIdCustom" Type="http://schemas.openxmlformats.org/'
                    b'officeDocument/2006/relationships/custom-properties" '
                    b'Target="docProps/custom.xml"/></Relationships>'))
            zout.writestr(item, data)
        zout.writestr("word/footnotes.xml", footnotes)
        zout.writestr("word/endnotes.xml", endnotes)
        zout.writestr("docProps/custom.xml", custom_props)

    return path


def all_xml_text(docx_path: str) -> str:
    """Every XML part of a .docx concatenated -- for asserting that a value
    is absent from the ENTIRE file, not just the parts a reader renders."""
    parts = []
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                parts.append(z.read(name).decode("utf-8", errors="replace"))
    return "\n".join(parts)
