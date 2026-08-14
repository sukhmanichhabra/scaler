"""
Unit tests for individual detectors and the end-to-end redaction pipeline.
Run with: python -m pytest tests/ -v
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

from app.core.detectors.base import DetectionConfig
from app.core.detectors.regex_detectors import (
    CreditCardDetector,
    EmailDetector,
    IPAddressDetector,
    PhoneDetector,
    SSNDetector,
)
from app.core.redactor import Redactor


class TestEmailDetector:
    def test_finds_standard_email(self):
        matches = EmailDetector().detect("Contact rashi.patil@gmail.com for details.")
        assert len(matches) == 1
        assert matches[0].text == "rashi.patil@gmail.com"
        assert matches[0].entity_type == "EMAIL"

    def test_ignores_non_email_text(self):
        matches = EmailDetector().detect("This costs 5.00@ per unit, not an email.")
        assert len(matches) == 0


class TestPhoneDetector:
    def test_finds_indian_mobile_with_country_code(self):
        matches = PhoneDetector().detect("Call +91 9876543210 for support.")
        assert len(matches) == 1
        assert matches[0].entity_type == "PHONE"

    def test_does_not_flag_short_reference_numbers(self):
        # "Order #45678" is the assignment's own example of a non-PII number
        matches = PhoneDetector().detect("Order #45678 was processed yesterday.")
        assert len(matches) == 0

    def test_does_not_flag_letter_prefixed_registration_numbers(self):
        # Found on a real prospectus: "SEBI registration no.: INM000013004"
        # was misread as a phone number because the digit run directly
        # follows a letter with no separator.
        matches = PhoneDetector().detect("SEBI registration no.: INM000013004")
        assert len(matches) == 0


class TestSSNDetector:
    def test_finds_valid_ssn_format(self):
        matches = SSNDetector().detect("Taxpayer ID on file: 219-09-9999.")
        assert len(matches) == 1

    def test_rejects_invalid_area_number(self):
        # area numbers 000, 666, and 900-999 are never issued
        matches = SSNDetector().detect("Reference code: 666-12-3456.")
        assert len(matches) == 0


class TestCreditCardDetector:
    def test_finds_valid_luhn_card(self):
        matches = CreditCardDetector().detect("Card number 4532015112830366 was charged.")
        assert len(matches) == 1

    def test_rejects_invalid_checksum(self):
        # same length/shape as a card number but fails Luhn
        matches = CreditCardDetector().detect("Tracking number 1234567890123456 was issued.")
        assert len(matches) == 0

    def test_does_not_swallow_trailing_text(self):
        matches = CreditCardDetector().detect("Card 4532015112830366 on file.")
        assert matches[0].text == "4532015112830366"
        assert matches[0].end == matches[0].start + len("4532015112830366")


class TestIPAddressDetector:
    def test_finds_ipv4(self):
        matches = IPAddressDetector().detect("Access logged from 192.168.1.104.")
        assert len(matches) == 1

    def test_rejects_out_of_range_octets(self):
        matches = IPAddressDetector().detect("Version number 999.999.999.999 is invalid.")
        assert len(matches) == 0


class TestOverlapResolutionAndPrecision:
    """Regression tests for the specific bugs found and fixed during development."""

    def test_ticket_and_order_numbers_survive(self):
        redactor = Redactor(config=DetectionConfig())
        text = "This draft carries internal Order #45678 and Ticket ID TCK-99881."
        result = redactor.redact(text)
        assert "Order #45678" in result.redacted_text
        assert "Ticket ID TCK-99881" in result.redacted_text

    def test_full_address_line_is_not_fragmented(self):
        redactor = Redactor(config=DetectionConfig())
        text = ("The registered office is situated at 4th Floor, Prestige Tech Park, "
                "Bellandur, Bangalore 560103, Karnataka, India.")
        result = redactor.redact(text)
        # none of the original address fragments should leak into the output
        assert "Prestige Tech Park" not in result.redacted_text
        assert "560103" not in result.redacted_text
        assert "Bellandur" not in result.redacted_text

    def test_consistent_replacement_across_document(self):
        redactor = Redactor(config=DetectionConfig())
        text = "Rohan Dey signed the form. Later, Rohan Dey confirmed by email."
        result = redactor.redact(text)
        fakes = [r["fake"] for r in result.replacements if r["original"] == "Rohan Dey"]
        assert len(fakes) == 2
        assert fakes[0] == fakes[1]

    def test_issuer_company_excluded_by_default(self):
        redactor = Redactor(config=DetectionConfig(), issuer_names={"novaweave technologies limited"})
        text = "Novaweave Technologies Limited today announced its results."
        result = redactor.redact(text)
        assert "Novaweave Technologies Limited" in result.redacted_text


class TestDefinedTermsStoplist:
    """A filing's own glossary is the highest-value precision signal there
    is -- see app/core/defined_terms.py."""

    def test_glossary_term_is_not_redacted_as_a_company(self):
        redactor = Redactor(config=DetectionConfig(),
                            defined_terms={"equity shares", "offer price"})
        text = "The Equity Shares will be allotted at the Offer Price determined by the Company."
        result = redactor.redact(text)
        assert "Equity Shares" in result.redacted_text
        assert "Offer Price" in result.redacted_text

    def test_derived_form_of_a_glossary_term_is_also_spared(self):
        # Filings write "pre-Offer" where the glossary only lists "Offer".
        redactor = Redactor(config=DetectionConfig(), defined_terms={"offer"})
        result = redactor.redact("Aggregate pre-Offer shareholding of our Promoters.")
        assert "pre-Offer" in result.redacted_text

    def test_real_company_is_still_redacted(self):
        # The stoplist must not become a blanket amnesty for capitalised text.
        redactor = Redactor(config=DetectionConfig(), defined_terms={"equity shares"})
        result = redactor.redact("The escrow account is maintained with HDFC Bank Limited.")
        assert "HDFC Bank Limited" not in result.redacted_text


class TestAddressContainingPhone:
    """Regression test for a real leak found on the source prospectus."""

    def test_address_survives_a_phone_number_inside_it(self):
        redactor = Redactor(config=DetectionConfig())
        text = ("The registered office is situated at 4th Floor, Prestige Tech Park, "
                "Bangalore 560103, Tel 020-45053237 India.")
        result = redactor.redact(text)
        # Previously the phone won the overlap and the whole address was
        # dropped, leaving the street address in cleartext.
        assert "Prestige Tech Park" not in result.redacted_text
        assert "560103" not in result.redacted_text
        assert "020-45053237" not in result.redacted_text


class TestContactBlockNames:
    """Contact blocks carry no sentence context, so NER alone misses them."""

    def test_name_after_contact_person_label(self):
        redactor = Redactor(config=DetectionConfig())
        text = "Telephone: +91 20 6769 4648 Contact Person: Manisha Shukla Website: www.example.com"
        result = redactor.redact(text)
        assert "Manisha Shukla" not in result.redacted_text

    def test_several_names_sharing_one_label(self):
        redactor = Redactor(config=DetectionConfig())
        text = "Contact Person: Eric Bacha/ Sachin Gawade/ Pravin Teli"
        result = redactor.redact(text)
        for name in ("Eric Bacha", "Sachin Gawade", "Pravin Teli"):
            assert name not in result.redacted_text

    def test_field_label_is_not_mistaken_for_a_name(self):
        redactor = Redactor(config=DetectionConfig())
        result = redactor.redact("Contact Person: Sharmila Joshi Email: a@b.com")
        assert "Email:" in result.redacted_text


class TestPublicBodiesAndCasing:
    def test_regulator_is_not_redacted(self):
        redactor = Redactor(config=DetectionConfig())
        text = "Registered with the Securities and Exchange Board of India under the regulations."
        result = redactor.redact(text)
        assert "Securities and Exchange Board of India" in result.redacted_text

    def test_all_caps_mention_stays_all_caps_and_matches_title_case(self):
        from app.core.faker_mapper import ConsistentFaker
        faker = ConsistentFaker()
        title = faker.fake_for("PERSON", "Kushal Subbayya Hegde")
        upper = faker.fake_for("PERSON", "KUSHAL SUBBAYYA HEGDE")
        assert upper.isupper()
        assert upper == title.upper()  # same identity, typography follows context


class TestNameVariantAliasing:
    def test_short_form_maps_to_the_same_fake_person(self):
        from app.core.detectors.registry import derive_name_aliases
        from app.core.faker_mapper import ConsistentFaker

        aliases = derive_name_aliases({"pushpa kushal hegde": "PERSON"})
        assert aliases == {"pushpa hegde": "pushpa kushal hegde"}

        faker = ConsistentFaker()
        faker.register_aliases(aliases)
        assert faker.fake_for("PERSON", "Pushpa Hegde") == \
            faker.fake_for("PERSON", "Pushpa Kushal Hegde")

    def test_ambiguous_short_form_is_not_aliased(self):
        from app.core.detectors.registry import derive_name_aliases
        # Two different people would collapse onto "rohit hegde".
        aliases = derive_name_aliases({
            "rohit kushal hegde": "PERSON",
            "rohit rajesh hegde": "PERSON",
        })
        assert "rohit hegde" not in aliases


class TestTableAwareDateOfBirth:
    def test_date_fires_when_the_column_header_says_date_of_birth(self):
        from app.core.detectors.regex_detectors import DateOfBirthDetector
        matches = DateOfBirthDetector().detect("14 March 1985", context="Date of Birth | Rohan Dey")
        assert len(matches) == 1

    def test_unrelated_date_in_a_table_is_left_alone(self):
        from app.core.detectors.regex_detectors import DateOfBirthDetector
        matches = DateOfBirthDetector().detect("12 January 2024", context="Board Meeting | Approval")
        assert len(matches) == 0


class TestDocxPipeline:
    """End-to-end .docx behaviour: merged cells, formatting, metadata."""

    @staticmethod
    def _build(tmp_path):
        from docx import Document
        doc = Document()
        doc.core_properties.author = "Real Author Name"
        p = doc.add_paragraph()
        p.add_run("Contact ")
        bold = p.add_run("Rohan Dey")
        bold.bold = True
        p.add_run(" on rohan.dey@gmail.com today.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = "Reach Priya Sharma at priya@example.com"
        path = str(tmp_path / "in.docx")
        doc.save(path)
        return path

    def test_formatting_metadata_and_merged_cells(self, tmp_path):
        from docx import Document
        from app.core.document_io import redact_file

        in_path = self._build(tmp_path)
        out_path = str(tmp_path / "out.docx")
        result = redact_file(in_path, out_path, Redactor(config=DetectionConfig()))
        out = Document(out_path)

        body = out.paragraphs[0]
        # Untouched runs keep their text and their formatting.
        assert body.runs[0].text == "Contact "
        assert any(r.bold for r in body.runs), "bold formatting was lost"
        assert "rohan.dey@gmail.com" not in body.text

        # Author metadata must not survive a redaction.
        assert not out.core_properties.author

        # A merged cell is redacted exactly once, not once per spanned column.
        merged_text = out.tables[0].cell(0, 0).text
        assert "priya@example.com" not in merged_text
        emails = [r for r in result.replacements
                  if r["type"] == "EMAIL" and r["original"] == "priya@example.com"]
        assert len(emails) == 1


class TestAddressCompleteness:
    """
    Regressions from the real prospectus, where addresses were not merely
    missed but SHREDDED: with no whole-span ADDRESS match, the ORG and
    PERSON detectors picked off the recognisable pieces and left the house
    number, PIN code and state sitting in cleartext between fake names.
    """

    @staticmethod
    def _addresses(text, context=""):
        from app.core.detectors.ner_detector import AddressDetector
        from app.core.detectors.registry import resolve_overlaps
        return [m.text for m in resolve_overlaps(AddressDetector().detect(text, context=context))]

    def test_two_addresses_in_one_sentence(self):
        # Neither address ends in punctuation within the old 180-char bound,
        # so the whole match used to fail and the first address leaked.
        text = ("KSH International Limited, having its Registered Office at 11/3, 11/4 and "
                "11/5, Village Birdewadi, Chakan Taluka-Khed, Pune - 410 501, Maharashtra, "
                "India and its Corporate Office at 201, Tower 2, Montreal Business Centre, "
                "Off Pallod Farms, Baner, Pune - 411 045, Maharashtra, India.")
        found = self._addresses(text)
        assert len(found) == 2
        assert "Village Birdewadi" in found[0] and "410 501" in found[0]
        assert "Montreal Business Centre" in found[1] and "411 045" in found[1]
        # the conjunction joining them must not be swallowed into the first
        assert "Corporate Office" not in found[0]

    def test_abbreviation_does_not_end_the_address(self):
        # "Plot No." used to truncate the span to "Plot No".
        found = self._addresses(
            "Our facility located at Plot No. J-25, Taloja Industrial Area, "
            "Village Padghe, Taluka Panvel, Raigad - 410 208, Maharashtra, India")
        assert len(found) == 1
        assert "410 208" in found[0]

    def test_address_labelled_by_its_table_column(self):
        # The "Registered Office" label lives in the header cell, not this one.
        found = self._addresses(
            "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune - 410 501\nMaharashtra, India",
            context="REGISTERED OFFICE")
        assert len(found) == 1 and "410 501" in found[0]

    def test_address_split_across_consecutive_lines(self):
        # Each line alone is too thin: one has the PIN, the other the keywords.
        assert self._addresses("Pune - 410 501",
                               context="11/3, Village Birdewadi | Maharashtra, India")
        assert self._addresses("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka-Khed",
                               context="KSH INTERNATIONAL LIMITED | Pune - 410 501")

    def test_prose_is_not_mistaken_for_an_address(self):
        for text, context in [
            ("Total borrowings are computed as current borrowings plus non-current borrowings.", ""),
            ("The details are as follows: John Smith, Chief Executive Officer", ""),
            ("Fiscal Year 2024 revenue grew 18%", "prose | Pune - 410 501"),
            ("The Board met on 12 January 2024.", "x | 410 501"),
            ("Equity Shares", "Offer | 411 045"),
        ]:
            assert self._addresses(text, context) == [], f"false positive on {text!r}"


class TestTextBoxContent:
    """
    Text boxes are invisible to python-docx's normal paragraph API, so PII
    inside one would otherwise survive redaction untouched.
    """

    @staticmethod
    def _doc_with_textbox(tmp_path):
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = Document()
        doc.add_paragraph("Ordinary body text.")
        # A legacy VML text box, which is how Word stores shape text.
        xml = (
            '<w:p %s xmlns:v="urn:schemas-microsoft-com:vml">'
            "<w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
            "<w:p><w:r><w:t>Reach Priya Sharma at priya.sharma@example.com</w:t></w:r></w:p>"
            "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>" % nsdecls("w")
        )
        doc.element.body.append(parse_xml(xml))
        path = str(tmp_path / "textbox.docx")
        doc.save(path)
        return path

    def test_pii_inside_a_text_box_is_redacted(self, tmp_path):
        import zipfile

        from app.core.document_io import redact_file

        in_path = self._doc_with_textbox(tmp_path)
        # Confirm the fixture really does hide the PII in a text box.
        with zipfile.ZipFile(in_path) as z:
            assert "txbxContent" in z.read("word/document.xml").decode("utf8")

        out_path = str(tmp_path / "out.docx")
        redact_file(in_path, out_path, Redactor(config=DetectionConfig()))

        with zipfile.ZipFile(out_path) as z:
            xml = z.read("word/document.xml").decode("utf8")
        assert "priya.sharma@example.com" not in xml
        assert "Ordinary body text." in xml  # untouched content survives


class TestIndianIdentifierDetectors:
    """
    The India-specific bonus types. Each pair below is deliberate: the
    positive case proves the detector fires, and the negative case proves
    the thing that keeps precision up -- either structural validation
    (IFSC's reserved '0', GSTIN's embedded PAN) or a required nearby
    keyword for the shapes too generic to trust alone.
    """

    @staticmethod
    def _texts(detector, text):
        return [m.text for m in detector.detect(text)]

    def test_ifsc_requires_the_reserved_zero(self):
        from app.core.detectors.regex_detectors import IFSCDetector
        det = IFSCDetector()
        assert self._texts(det, "Transfer via IFSC HDFC0001234 today.") == ["HDFC0001234"]
        # 5th character is not '0' -- not an IFSC, and the 1-in-36 chance of
        # a random string landing one there is what makes this a real check.
        assert self._texts(det, "Reference code ABCD1234567 in the ledger.") == []

    def test_gstin_requires_a_well_formed_embedded_pan(self):
        from app.core.detectors.regex_detectors import GSTINDetector
        det = GSTINDetector()
        assert self._texts(det, "GSTIN: 27AAAPL1234C1Z5 on the invoice.") == ["27AAAPL1234C1Z5"]
        # Same length, but the embedded PAN section is malformed.
        assert self._texts(det, "Code 271234AAAAC1Z5X here.") == []

    def test_upi_id_matches_only_known_psp_handles(self):
        from app.core.detectors.regex_detectors import UPIIdDetector
        det = UPIIdDetector()
        found = self._texts(det, "Pay rohan.dey@okhdfcbank or 9876543210@ybl before noon.")
        assert found == ["rohan.dey@okhdfcbank", "9876543210@ybl"]
        # An arbitrary "word@word" is not a UPI ID -- matching those would
        # flag ordinary prose containing an @ sign.
        assert self._texts(det, "Let's meet@noon in the boardroom.") == []

    def test_upi_id_does_not_swallow_ordinary_emails(self):
        # EmailDetector owns anything with a dotted TLD; the two patterns
        # must not both claim the same span.
        from app.core.detectors.regex_detectors import UPIIdDetector
        assert self._texts(UPIIdDetector(), "Write to rohan.dey@gmail.com please.") == []

    @pytest.mark.parametrize("detector_name,positive,negative,keyword_text", [
        ("PassportNumberDetector", "M1234567",
         "Reference M1234567 in the ledger.", "Passport No: M1234567 issued."),
        ("VoterIdDetector", "ABC1234567",
         "Batch ABC1234567 dispatched.", "Voter ID: ABC1234567 verified."),
        ("BankAccountNumberDetector", "123456789012",
         "Order number 123456789012 shipped.", "Account No: 123456789012 at the bank."),
    ])
    def test_generic_shapes_require_a_nearby_keyword(self, detector_name, positive,
                                                     negative, keyword_text):
        import app.core.detectors.regex_detectors as rd
        det = getattr(rd, detector_name)()
        assert positive in self._texts(det, keyword_text)
        # Without the label these shapes are indistinguishable from the
        # reference codes a formal document is full of.
        assert self._texts(det, negative) == []

    def test_driving_license(self):
        from app.core.detectors.regex_detectors import DrivingLicenseDetector
        det = DrivingLicenseDetector()
        assert self._texts(det, "Driving Licence No: MH12 20110012345 valid.") == \
            ["MH12 20110012345"]
        assert self._texts(det, "Batch MH12 20110012345 was processed.") == []

    def test_new_types_are_registered_end_to_end(self):
        # A detector that exists but isn't wired into the registry, the
        # config default, or the faker would silently never run.
        from app.core.detectors.base import ALL_PII_TYPES
        from app.core.detectors.registry import _TYPE_TO_DETECTOR
        new_types = {"IFSC_CODE", "GSTIN", "UPI_ID", "PASSPORT_NUMBER",
                     "VOTER_ID", "DRIVING_LICENSE", "BANK_ACCOUNT_NUMBER"}
        assert new_types <= set(ALL_PII_TYPES)
        assert new_types <= set(_TYPE_TO_DETECTOR)
        assert new_types <= DetectionConfig().enabled_types

    def test_each_new_type_gets_a_realistic_fake(self):
        from app.core.faker_mapper import ConsistentFaker
        faker = ConsistentFaker()
        cases = {
            "IFSC_CODE": "HDFC0001234",
            "GSTIN": "27AAAPL1234C1Z5",
            "UPI_ID": "rohan.dey@okhdfcbank",
            "PASSPORT_NUMBER": "M1234567",
            "VOTER_ID": "ABC1234567",
            "DRIVING_LICENSE": "MH12 20110012345",
            "BANK_ACCOUNT_NUMBER": "123456789012",
        }
        for entity_type, original in cases.items():
            fake = faker.fake_for(entity_type, original)
            assert fake and fake != original
            # No type should fall through to the generic hash placeholder.
            assert not fake.startswith("REDACTED-"), f"{entity_type} has no generator"
        # Account-number length is itself a mild signal, so it's preserved.
        assert len(faker.fake_for("BANK_ACCOUNT_NUMBER", "123456789012")) == 12

    def test_redaction_is_consistent_for_new_types(self):
        redactor = Redactor(config=DetectionConfig())
        text = ("Account No: 123456789012 with IFSC HDFC0001234. "
                "Confirm to Account No: 123456789012 again.")
        result = redactor.redact(text)
        account_fakes = [r["fake"] for r in result.replacements
                         if r["type"] == "BANK_ACCOUNT_NUMBER"]
        assert len(account_fakes) == 2 and account_fakes[0] == account_fakes[1]
        assert "123456789012" not in result.redacted_text
        assert "HDFC0001234" not in result.redacted_text


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
