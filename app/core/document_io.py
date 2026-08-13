"""
Handles reading supported input formats (.docx, .pdf, .txt) and always
writing the redacted result out as a .docx (per the assignment's output
requirement).

For .docx input we redact run-by-run so that existing formatting (bold,
italics, fonts) survives -- when a detected PII span sits entirely inside
one run we edit that run's text in place; when it happens to straddle a run
boundary we fall back to concatenating into the first run (rare, e.g. a name
that's partly bold). For .pdf / .txt input there is no original formatting
to preserve, so we generate a clean new .docx from the extracted text.
"""
from __future__ import annotations

import os
from dataclasses import dataclass
from typing import Dict, List, Optional

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .defined_terms import extract_defined_terms
from .detectors.registry import (apply_known_names, build_known_names,
                                 derive_name_aliases, derive_org_short_forms)
from .docx_hidden_content import (HiddenContentReport, detect_unredactable_content,
                                  redact_comments, redact_custom_properties)
from .docx_revisions import flatten_revisions
from .redactor import Redactor, RedactionResult


@dataclass(frozen=True)
class _Replacement:
    """A resolved (span -> fake value) edit, ready to write into runs."""

    start: int
    end: int
    fake: str


def _is_heading(paragraph: Paragraph) -> bool:
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        return False
    return style_name.startswith(("Heading", "Title"))


def _apply_matches_to_runs(paragraph: Paragraph, matches: List) -> None:
    """
    Writes replacements back into the paragraph's runs, preserving each
    run's own formatting.

    A run is the unit of formatting in a .docx, and a paragraph is usually
    many runs ("Contact " / "Rohan Dey" / " on +91..."). Rewriting only the
    runs a match actually touches keeps bold/italic/font on every untouched
    run -- which matters a lot in a formal filing where nearly every
    paragraph carries mixed emphasis.

    When a match straddles a run boundary the replacement text is written
    into the first run it touches and the covered part of the later runs is
    removed, so content is always correct; only that one span adopts the
    first run's style.
    """
    runs = paragraph.runs
    spans = []  # (run, start, end) offsets into the concatenated text
    cursor = 0
    for run in runs:
        spans.append((run, cursor, cursor + len(run.text)))
        cursor += len(run.text)

    # Apply back-to-front so earlier offsets stay valid.
    for m in sorted(matches, key=lambda x: -x.start):
        replacement = m.fake
        for run, r_start, r_end in spans:
            if r_end <= m.start or r_start >= m.end:
                continue  # run untouched by this match
            local_start = max(m.start - r_start, 0)
            local_end = min(m.end - r_start, r_end - r_start)
            head = run.text[:local_start]
            tail = run.text[local_end:]
            run.text = head + replacement + tail
            replacement = ""  # only the first covered run receives the text


def _redact_paragraph_in_place(paragraph: Paragraph, redactor: Redactor, all_replacements: list, all_matches: list,
                                known_names: Optional[Dict[str, str]] = None,
                                matches: Optional[List] = None):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return
    if matches is None:
        matches = redactor.detect(full_text, skip_ner=_is_heading(paragraph), known_names=known_names)
    result = redactor.redact_matches(full_text, matches)
    all_replacements.extend(result.replacements)
    all_matches.extend(result.matches)
    if not result.matches:
        return  # nothing changed, don't touch runs / formatting

    # Pair each match with the fake value chosen for it, then splice those
    # into the runs. `replacements` is in document order; matches are too.
    ordered = sorted(result.matches, key=lambda x: x.start)
    annotated = []
    for m, rep in zip(ordered, result.replacements):
        annotated.append(_Replacement(m.start, m.end, rep["fake"]))
    _apply_matches_to_runs(paragraph, annotated)


# A section can carry three independent header/footer variants (default,
# distinct first page, distinct even pages). Reading only `.header`/`.footer`
# silently skips the other two, which is exactly where a letterhead address
# or a running contact line tends to live.
_HEADER_FOOTER_ATTRS = (
    "header", "footer",
    "first_page_header", "first_page_footer",
    "even_page_header", "even_page_footer",
)


def _textbox_paragraphs(element, parent) -> List[Paragraph]:
    """
    Paragraphs inside text boxes and shapes.

    `python-docx` exposes only paragraphs that are direct children of the
    body (or of a cell/header), so anything Word puts inside a text box --
    `<w:txbxContent>`, reached via a drawing or a legacy VML shape -- is
    invisible to the normal API and would pass through a text-only
    redaction completely untouched. Cover pages and letterheads are exactly
    where designers reach for a text box, so this is worth walking
    explicitly rather than hoping no document uses one.

    Note: Word often stores a shape twice inside `mc:AlternateContent` (a
    modern `mc:Choice` and a legacy `mc:Fallback` carrying the same text).
    Both copies are redacted, because either one may be what a given reader
    renders. They resolve to identical fake values via the replacement
    cache, so the document stays consistent; the only side effect is that
    such a span can appear twice in the audit log.
    """
    found: List[Paragraph] = []
    for textbox in element.iter(qn("w:txbxContent")):
        for p_element in textbox.iter(qn("w:p")):
            found.append(Paragraph(p_element, parent))
    return found


def _note_part(document: Document, reltype_suffix: str):
    for rel in document.part.rels.values():
        if rel.reltype.endswith(reltype_suffix) and not rel.is_external:
            return rel.target_part
    return None


def _footnote_and_endnote_paragraphs(document: Document) -> "tuple[List[Paragraph], list]":
    """
    Parses footnotes.xml / endnotes.xml (if present) and returns their
    paragraphs alongside a list of `(part, root_element)` pairs still to be
    written back after redaction.

    Neither part has a `python-docx` object model (see
    `docx_hidden_content.py`), so this parses the raw XML once here, hands
    back live `Paragraph` objects wrapping elements from that SAME parsed
    tree, and leaves the caller responsible for re-serializing each root
    into its part's `_blob` after mutating those paragraphs -- mirroring
    exactly how `_redact_note_part` in `docx_hidden_content.py` does it for
    a single detect-and-redact-immediately call, but split across two steps
    here so footnotes/endnotes can share the SAME detect-once /
    known-names pass as the rest of the document instead of being redacted
    in isolation.
    """
    from docx.oxml import parse_xml

    paragraphs: List[Paragraph] = []
    pending_writeback = []
    for reltype_suffix, container_tag in (("/footnotes", "w:footnote"), ("/endnotes", "w:endnote")):
        part = _note_part(document, reltype_suffix)
        if part is None:
            continue
        root = parse_xml(part.blob)
        flatten_revisions(root)
        for note in root.iter(qn(container_tag)):
            if note.get(qn("w:type")) in ("separator", "continuationSeparator"):
                continue
            for p_element in note.findall(qn("w:p")):
                paragraphs.append(Paragraph(p_element, part))
        pending_writeback.append((part, root))
    return paragraphs, pending_writeback


def _write_back_notes(pending_writeback: list) -> None:
    from lxml import etree
    for part, root in pending_writeback:
        part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _iter_docx_paragraphs(doc: Document) -> List[tuple]:
    """
    Every paragraph the redaction pipeline touches, paired with its
    structural context: body text, table cells (deduped against
    merged-cell duplication -- see the note below), and every
    header/footer variant of every section.

    Context is the cell's column header and row label, which is how a
    "Date of Birth" column gets associated with the dates underneath it.
    Body paragraphs carry no context and get an empty string.
    """
    # Body paragraphs carry their immediate neighbours as context. A postal
    # address on a cover page is routinely typeset one line per paragraph --
    # "11/3, Village Birdewadi" / "Pune - 410 501" / "Maharashtra, India" --
    # and no single line then holds enough signal to be recognised on its
    # own, so the address escapes redaction and gets shredded into fake
    # names instead. The neighbouring lines supply the missing evidence.
    body = list(doc.paragraphs)
    body_texts = [p.text.strip() for p in body]
    paragraphs: List[tuple] = []
    for i, paragraph in enumerate(body):
        previous_text = body_texts[i - 1] if i > 0 else ""
        next_text = body_texts[i + 1] if i + 1 < len(body_texts) else ""
        neighbours = " | ".join(t for t in (previous_text, next_text) if t)[:220]
        paragraphs.append((paragraph, neighbours))
    paragraphs.extend((p, "") for p in _textbox_paragraphs(doc.element.body, doc))
    for table in doc.tables:
        # Horizontally/vertically merged cells make python-docx yield the
        # SAME underlying <w:tc> element multiple times from `row.cells`
        # (once per spanned grid column/row). Without deduping, a merged
        # cell gets redacted more than once -- the second pass detects PII
        # in the already-redacted (fake) text from the first pass and
        # redacts that too, corrupting the output and polluting the audit
        # log with fake values mislabeled as "original". Track which cells
        # we've already processed per table to redact each physical cell
        # exactly once.
        seen_cells = set()
        header_cells = []
        try:
            header_cells = [c.text.strip() for c in table.rows[0].cells]
        except (IndexError, AttributeError):  # pragma: no cover - malformed table
            header_cells = []
        for row_idx, row in enumerate(table.rows):
            try:
                row_label = row.cells[0].text.strip()
            except (IndexError, AttributeError):  # pragma: no cover
                row_label = ""
            for col_idx, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                column_header = header_cells[col_idx] if col_idx < len(header_cells) else ""
                # Don't feed a cell its own text back as its context.
                parts = [t for t in (column_header, row_label if row_idx else "")
                         if t and t != cell.text.strip()]
                context = " | ".join(parts)[:200]
                paragraphs.extend((p, context) for p in cell.paragraphs)
    seen_parts = set()
    for section in doc.sections:
        for attr in _HEADER_FOOTER_ATTRS:
            part = getattr(section, attr, None)
            if part is None:
                continue
            # Sections routinely inherit the same header/footer part; only
            # walk each physical part once.
            key = id(part._element) if hasattr(part, "_element") else id(part)
            if key in seen_parts:
                continue
            seen_parts.add(key)
            if hasattr(part, "_element"):
                flatten_revisions(part._element)
            paragraphs.extend((p, "") for p in part.paragraphs)
            for table in getattr(part, "tables", []):
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend((p, "") for p in cell.paragraphs)
            if hasattr(part, "_element"):
                paragraphs.extend((p, "") for p in _textbox_paragraphs(part._element, part))
    return paragraphs


def _org_alias_pairs(known_names: Dict[str, str], short_forms: Dict[str, str]):
    """Pairs each derived company short form with the full name it came from,
    so both resolve to a single fake company rather than two."""
    for short in short_forms:
        for full_name, entity_type in known_names.items():
            if entity_type == "ORG" and full_name.startswith(short + " "):
                yield short, full_name
                break


def _scrub_metadata(doc: Document) -> None:
    """
    Clears document properties that carry identity information.

    A .docx stores the author, the last person to save it, and often the
    authoring company in `docProps/core.xml` -- none of which is visible in
    the page text, all of which survives a text-only redaction. Leaving
    them intact is a classic way for a "redacted" document to still name
    the people behind it.
    """
    props = doc.core_properties
    for field in ("author", "last_modified_by", "category", "comments",
                  "content_status", "identifier", "keywords", "subject", "title"):
        try:
            setattr(props, field, "")
        except (AttributeError, ValueError):  # pragma: no cover - property is read-only
            continue


def redact_docx(input_path: str, output_path: str, redactor: Redactor) -> RedactionResult:
    doc = Document(input_path)

    # Tracked-change insertions and hyperlink display text are invisible to
    # `paragraph.runs` until unwrapped -- see `docx_revisions.py` for the
    # verified failure mode (an insertion-only paragraph reports 0 runs).
    # Must run before anything reads paragraph text, body-wide, since a
    # revision can appear in the body, a table cell, or a text box alike.
    flatten_revisions(doc.element.body)

    paragraphs = _iter_docx_paragraphs(doc)  # also flattens each header/footer part

    # Footnotes and endnotes have no `python-docx` object model at all (see
    # `docx_hidden_content.py`), so their paragraphs are parsed separately
    # and folded into the SAME detect-once pass as everything else, so a
    # name only mentioned in a footnote still contributes to -- and
    # benefits from -- the document-wide consistency sweep below.
    note_paragraphs, note_writeback = _footnote_and_endnote_paragraphs(doc)
    paragraphs.extend((p, "") for p in note_paragraphs)

    # A prospectus defines its own vocabulary ("Equity Shares", "the Offer
    # Price"); those defined terms are emphatically NOT PII, and feeding
    # them to the NER detectors as a stoplist is the single biggest
    # precision win available here. See `app/core/defined_terms.py`.
    if not redactor.defined_terms:
        redactor.defined_terms = extract_defined_terms(doc)

    # Detect ONCE per paragraph, keeping the matches. The old shape ran a
    # detect-only pass and then re-detected during replacement, paying for
    # the (dominant) NER cost twice over the whole document.
    detected: List[tuple] = []
    for p, context in paragraphs:
        text = "".join(run.text for run in p.runs)
        if not text.strip():
            continue
        detected.append((p, text, redactor.detect(text, skip_ner=_is_heading(p), context=context)))

    # With every mention in hand, build the registry of names confidently
    # seen somewhere in the document, so a casing variant NER missed on one
    # specific occurrence (e.g. an ALL-CAPS cover-page declaration) still
    # gets caught. This refinement is regex-only -- no model involved.
    known_names = build_known_names([m for _, _, ms in detected for m in ms])

    # Shortened spellings of a known person ("Pushpa Hegde" for "Pushpa
    # Kushal Hegde") are both swept in as matches AND pointed at the full
    # name's fake identity, so the same person never becomes two fake
    # people -- and a short-form mention NER missed doesn't leak.
    aliases = derive_name_aliases(known_names)
    # The same treatment for companies: "ICICI Securities Limited" is also
    # written "ICICI Securities", and the short form leaks wherever NER
    # happens to miss it.
    org_short_forms = derive_org_short_forms(known_names)
    aliases.update({short: full for short, full in _org_alias_pairs(known_names, org_short_forms)})
    if aliases:
        redactor.faker.register_aliases(aliases)
    if org_short_forms or aliases:
        known_names = {
            **known_names,
            **{short: "PERSON" for short in derive_name_aliases(known_names)},
            **org_short_forms,
        }

    all_replacements: List[dict] = []
    all_matches = []
    for p, text, matches in detected:
        final = apply_known_names(text, matches, known_names)
        _redact_paragraph_in_place(p, redactor, all_replacements, all_matches, matches=final)

    # Footnote/endnote paragraphs were mutated in place above (they're the
    # SAME `Paragraph` objects `detected` iterated over), but that mutation
    # only touched the parsed-in-memory tree -- `doc.save()` has no idea
    # that tree exists unless it's written back into the part's blob first.
    _write_back_notes(note_writeback)

    # Comments get their own detect-and-redact pass rather than sharing the
    # main `detected` list: they're independent review annotations, not
    # part of the document's substantive content, so the extra plumbing to
    # fold them into the SAME known-names sweep wasn't judged worth it --
    # each comment is still fully and correctly redacted on its own merits,
    # it just doesn't cross-pollinate consistency with the body. The
    # comment AUTHOR (a real name in the metadata, independent of what the
    # comment text says) is always scrubbed regardless.
    def _redact_comment_paragraph(paragraph: Paragraph) -> bool:
        text = "".join(run.text for run in paragraph.runs)
        if not text.strip():
            return False
        matches = redactor.detect(text)
        if not matches:
            return False
        result = redactor.redact_matches(text, matches)
        all_replacements.extend(result.replacements)
        all_matches.extend(result.matches)
        ordered = sorted(result.matches, key=lambda x: x.start)
        annotated = [_Replacement(m.start, m.end, rep["fake"])
                     for m, rep in zip(ordered, result.replacements)]
        _apply_matches_to_runs(paragraph, annotated)
        return True

    def _fake_comment_author(name: str) -> str:
        fake_name = redactor.faker.fake_for("PERSON", name)
        # Recorded as a replacement for two reasons: it completes the audit
        # trail (a comment author is redacted PII exactly as much as a name
        # in the body), and it registers the fake name as a KNOWN fake value
        # -- without this, the residual scan's re-detection pass has no way
        # to tell "Jane Reviewer" the correctly-substituted fake author from
        # a leaked real one, and reports a false positive on its own output.
        all_replacements.append({
            "type": "PERSON", "original": name, "fake": fake_name,
            "confidence": 1.0, "source": "metadata",
        })
        return fake_name

    redact_comments(doc, _redact_comment_paragraph, _fake_comment_author)

    warnings: List[str] = []
    images, objects = detect_unredactable_content(doc)
    if images:
        warnings.append(
            f"{images} embedded image(s) detected. Text inside an image (a scanned "
            f"signature, a photographed ID card, a screenshot) is NOT scanned -- this "
            f"tool does not perform OCR. Review images manually before sharing."
        )
    if objects:
        warnings.append(
            f"{objects} embedded object(s) detected (e.g. an embedded spreadsheet or "
            f"file). Embedded object content is NOT scanned. Review manually."
        )

    _scrub_metadata(doc)
    doc.save(output_path)

    # Custom document properties (docProps/custom.xml) have no python-docx
    # object model, so this is a raw zip rewrite of the file that was just
    # saved -- see docx_hidden_content.py for why.
    def _redact_property_text(text: str) -> str:
        matches = redactor.detect(text)
        if not matches:
            return text
        result = redactor.redact_matches(text, matches)
        all_replacements.extend(result.replacements)
        all_matches.extend(result.matches)
        return result.redacted_text

    custom_props_redacted = redact_custom_properties(output_path, _redact_property_text)
    if custom_props_redacted:
        warnings.append(
            f"{custom_props_redacted} custom document propert(y/ies) contained PII and were redacted."
        )

    combined_text = "\n".join(p.text for p in doc.paragraphs)
    return RedactionResult(redacted_text=combined_text, matches=all_matches,
                           replacements=all_replacements, warnings=warnings)


def redact_text_to_docx(input_text: str, output_path: str, redactor: Redactor,
                         title: str = "Redacted Document") -> RedactionResult:
    # Detect once, then apply the name-consistency refinement over those
    # same matches (see `redact_docx` for why this ordering matters).
    matches = redactor.detect(input_text)
    known_names = build_known_names(matches)
    result = redactor.redact_matches(input_text, apply_known_names(input_text, matches, known_names))
    doc = Document()
    doc.add_heading(title, level=1)
    for para in result.redacted_text.split("\n"):
        doc.add_paragraph(para if para.strip() else "")
    doc.save(output_path)
    return result


def extract_pdf_text(input_path: str) -> str:
    chunks = []
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def redact_file(input_path: str, output_path: str, redactor: Redactor,
                verify: bool = True) -> RedactionResult:
    """
    Dispatches on file extension. Always produces a .docx at output_path.

    When `verify` is True (the default), the finished file is re-opened and
    re-scanned for PII before this returns -- see `verification.py` for
    what that does and doesn't catch. This roughly doubles wall-clock time
    on a large document (one extra full-document detection pass), which is
    the deliberate trade being made: catching a leak before it leaves this
    function is worth more than the extra seconds. Set `verify=False` to
    skip it (e.g. for the evaluation harnesses, which score raw detection
    against hand-labelled ground truth and have no "finished file" to
    re-check).
    """
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        result = redact_docx(input_path, output_path, redactor)
    elif ext == ".pdf":
        text = extract_pdf_text(input_path)
        result = redact_text_to_docx(text, output_path, redactor,
                                     title=os.path.basename(input_path))
    elif ext == ".txt":
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        result = redact_text_to_docx(text, output_path, redactor,
                                     title=os.path.basename(input_path))
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx, .pdf, .txt")

    if verify:
        from .verification import verify_docx
        result.residual = verify_docx(output_path, result.replacements, redactor)

    return result
